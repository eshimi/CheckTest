from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def _mask_id(s):
    """ログインIDの中間部分をアスタリスクで伏字にする（先頭3文字＋末尾2文字を残す）"""
    s = str(s)
    if len(s) <= 5:
        return s[0] + '*' * (len(s) - 1)
    return s[:3] + '*' * (len(s) - 5) + s[-2:]


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def rgb(hex_color):
    h = hex_color.lstrip("#")
    return RGBColor(*bytes.fromhex(h))


def add_section_title(doc, text, primary):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = rgb(primary)
    # 背景色（shading on paragraph）
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEF2F7")
    pPr.append(shd)
    return p


SCENE_ORDER = ["シーン1", "シーン2", "シーン3"]


def generate_word_report(examinee: dict, output_path: str,
                         scene_scores: dict = None, overall_comment: str = "",
                         comp_comments: dict = None):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    PRIMARY = "1A3A5C"
    comp_cols = examinee.get("competency_cols", [])
    comp_comments = comp_comments or {}

    # ── タイトル ──────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("認定プログラムフィードバックシート")
    r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = rgb(PRIMARY)

    # ── 受験者情報（1行） ──────────────────────────────────────────────────────
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    masked = _mask_id(examinee.get("id", examinee.get("name", "")))
    sub_run = sub.add_run(
        f"{masked} 様　｜　部門: {examinee.get('department', '')}　｜　"
        f"班: {examinee.get('genre', '未設定')}　｜　"
        f"総合得点率: {examinee['percentage']}%"
    )
    sub_run.font.size = Pt(11)
    doc.add_paragraph()

    # ── 総評・改善提案を分割 ────────────────────────────────────────────────────
    comment_parts = overall_comment.split("\n\n") if overall_comment else [""]

    # ── Section 1: 総評 ────────────────────────────────────────────────────────
    add_section_title(doc, "1. 総評", PRIMARY)
    doc.add_paragraph(comment_parts[0] if comment_parts else "")

    # ── Section 2: コンピテンシーごとの評価 ────────────────────────────────────
    add_section_title(doc, "2. コンピテンシーごとの評価", PRIMARY)

    if comp_cols:
        tbl = doc.add_table(rows=1 + len(comp_cols), cols=3)
        tbl.style = "Table Grid"
        for hdr, txt in zip(tbl.rows[0].cells, ["コンピテンシー", "得点率", "評価コメント"]):
            hdr.text = txt
            set_cell_bg(hdr, PRIMARY)
            for run in hdr.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
        for i, c in enumerate(comp_cols):
            row = tbl.rows[i + 1]
            rate = examinee["comp_rates"].get(c, 0)
            row.cells[0].text = c
            row.cells[1].text = f"{rate}%"
            row.cells[2].text = comp_comments.get(c, "")
            bg = "D4EDDA" if rate >= 75 else ("FFF3CD" if rate >= 50 else "F8D7DA")
            set_cell_bg(row.cells[1], bg)
            for run in row.cells[0].paragraphs[0].runs:
                run.font.bold = True

    doc.add_paragraph()

    # ── Section 3: 改善提案 ────────────────────────────────────────────────────
    add_section_title(doc, "3. 改善提案", PRIMARY)
    doc.add_paragraph(comment_parts[1] if len(comment_parts) > 1 else "")
    doc.add_paragraph()

    # ── シーン別得点 ────────────────────────────────────────────────────────────
    if scene_scores:
        h2 = doc.add_heading("シーン別得点", level=2)
        for r in h2.runs:
            r.font.color.rgb = rgb(PRIMARY)

        ordered_scenes = [s for s in SCENE_ORDER if s in scene_scores] + \
                         [s for s in scene_scores if s not in SCENE_ORDER]
        for s in ordered_scenes:
            ss = scene_scores[s]
            pct = ss["pct"]
            bg_scene = "D4EDDA" if pct >= 75 else ("FFF3CD" if pct >= 50 else "F8D7DA")

            scene_tbl = doc.add_table(rows=1, cols=2)
            scene_tbl.style = "Table Grid"
            scene_tbl.rows[0].cells[0].text = s
            scene_tbl.rows[0].cells[1].text = f"{ss['total']} / {ss['max']}点　({pct}%)"
            set_cell_bg(scene_tbl.rows[0].cells[0], PRIMARY)
            set_cell_bg(scene_tbl.rows[0].cells[1], bg_scene)
            for run in scene_tbl.rows[0].cells[0].paragraphs[0].runs:
                run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for run in scene_tbl.rows[0].cells[1].paragraphs[0].runs:
                run.font.bold = True

            comp_bd = ss.get("comp_breakdown", {})
            if comp_bd:
                comp_tbl = doc.add_table(rows=1 + len(comp_bd), cols=3)
                comp_tbl.style = "Table Grid"
                for hdr, txt in zip(comp_tbl.rows[0].cells, ["コンピテンシー", "得点", "達成率"]):
                    hdr.text = txt; set_cell_bg(hdr, "E8EEF5")
                    for run in hdr.paragraphs[0].runs:
                        run.font.bold = True
                for i, (c, cb) in enumerate(comp_bd.items()):
                    row = comp_tbl.rows[i + 1]
                    row.cells[0].text = c
                    row.cells[1].text = f"{cb['total']} / {cb['max']}点"
                    row.cells[2].text = f"{cb['pct']}%"
                    bg = "D4EDDA" if cb["pct"] >= 75 else ("FFF3CD" if cb["pct"] >= 50 else "F8D7DA")
                    set_cell_bg(row.cells[2], bg)
            doc.add_paragraph()

    # ── 問題別採点結果 ──────────────────────────────────────────────────────────
    h2 = doc.add_heading("問題別採点結果", level=2)
    for r in h2.runs:
        r.font.color.rgb = rgb(PRIMARY)

    answers_by_scene = {}
    for ans in examinee.get("answers", []):
        s = ans.get("question_genre", "その他")
        answers_by_scene.setdefault(s, []).append(ans)
    ordered_scenes = [s for s in SCENE_ORDER if s in answers_by_scene] + \
                     [s for s in answers_by_scene if s not in SCENE_ORDER]

    q_num = 0
    for scene in ordered_scenes:
        scene_ans_list = answers_by_scene[scene]
        s_total = sum(sum(a.get("competency_scores", {}).values()) for a in scene_ans_list)
        s_max = sum(len(a.get("competency_scores", {})) * 3 for a in scene_ans_list)
        s_pct = round(s_total / s_max * 100, 1) if s_max > 0 else 0

        sh = doc.add_heading(f"{scene}　{s_total}/{s_max}点 ({s_pct}%)", level=3)
        for r in sh.runs:
            r.font.color.rgb = rgb(PRIMARY)
        sh.paragraph_format.space_before = Pt(6)

        for ans in scene_ans_list:
            q_num += 1
            comp_scores = ans.get("competency_scores", {})
            total_q = sum(comp_scores.values())
            max_q = len(comp_scores) * 3

            # 問タイトル
            qp = doc.add_paragraph()
            qr = qp.add_run(f"問{q_num}【{ans['question_type']}】　{ans['question_text'][:80]}")
            qr.font.bold = True; qr.font.size = Pt(10)

            # 強制0点
            if ans.get("forced_zero"):
                fp = doc.add_paragraph()
                fr = fp.add_run(f"⚠ 強制0点適用：{ans.get('forced_zero_reason', '')}")
                fr.font.color.rgb = RGBColor(0x8C, 0x2A, 0x1A)
                fr.font.size = Pt(9)

            # コンピテンシー別得点
            score_str = "　".join(f"{c}: {s}/3" for c, s in comp_scores.items())
            score_p = doc.add_paragraph()
            score_r = score_p.add_run(f"得点 {total_q}/{max_q}点　│　{score_str}")
            score_r.font.size = Pt(9)
            score_r.font.bold = True

            # 受験者の回答
            rows_data = [
                ("受験者の回答", ans.get("answer_text") or "（無回答）"),
            ]

            # competency_reasons
            comp_reasons = ans.get("competency_reasons", {})
            if comp_reasons:
                reasons_text = "\n".join(f"{c}：{r}" for c, r in comp_reasons.items())
                rows_data.append(("採点理由", reasons_text))

            tbl = doc.add_table(rows=len(rows_data), cols=2)
            tbl.style = "Table Grid"
            for r_idx, (label, val) in enumerate(rows_data):
                row = tbl.rows[r_idx]
                row.cells[0].text = label
                row.cells[1].text = str(val)
                set_cell_bg(row.cells[0], "E8EEF5")
                for run in row.cells[0].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(9)

            doc.add_paragraph()

    doc.save(output_path)
