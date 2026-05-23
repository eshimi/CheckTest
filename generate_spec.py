# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
WHITE   = RGBColor(0xff, 0xff, 0xff)
GRAY    = RGBColor(0x88, 0x88, 0x88)

def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = 'Yu Gothic'  # ascii + hAnsi をAPIで設定
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 日本語フォント（eastAsia）を強制設定、テーマ上書きを除去
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    # テーマフォント属性を除去（これが明示指定を上書きする原因）
    for attr in [qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme'), qn('w:cstheme')]:
        rFonts.attrib.pop(attr, None)
    rFonts.set(qn('w:ascii'),    'Yu Gothic')
    rFonts.set(qn('w:hAnsi'),   'Yu Gothic')
    rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    rFonts.set(qn('w:cs'),      'Yu Gothic')

def add_border(p, color_hex='1a3a5c'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=14, bold=True, color=PRIMARY)
    add_border(p)

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run('■ ' + text)
    set_font(r, size=11.5, bold=True, color=PRIMARY)

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    set_font(r, size=10.5)

def bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run('・ ' + text)
    set_font(r, size=10.5)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    set_font(r, size=10, bold=True, color=PRIMARY)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # ヘッダー行
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, '1a3a5c')
    # データ行
    for ri, row in enumerate(rows):
        bg = 'f0f4f8' if ri % 2 == 0 else 'ffffff'
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val))
            set_font(r, size=10)
            set_cell_bg(cell, bg)
    # 列幅
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ════════════════════════════════════════
# タイトル
# ════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('中上級者認定テスト　採点AIシステム　仕様書')
set_font(r, size=16, bold=True, color=PRIMARY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(16)
r2 = p2.add_run(
    '作成日：' + datetime.date.today().strftime('%Y年%m月%d日') +
    '　　Powered by Claude AI（Anthropic）'
)
set_font(r2, size=9, color=GRAY)

doc.add_paragraph()

# ════════════════════════════════════════
# 1. 試験の概要
# ════════════════════════════════════════
heading1('1. 試験の概要')
add_table(
    ['項目', '内容'],
    [
        ['試験名',     '中上級者認定テスト'],
        ['目的',       '災害対応シナリオへの対処能力を測定し、受験者の成長を支援する'],
        ['形式',       '記述式（シナリオに対する自由回答）'],
        ['問題数',     '全22問（シーン1〜3に分かれて出題）'],
        ['採点方式',   'Claude AI（Anthropic）によるルーブリック自動採点'],
        ['受験者単位', '班（ジャンル）ごとに専用の問題・ルーブリックを使用'],
    ],
    col_widths=[4, 11.5]
)
doc.add_paragraph()
heading2('シーン構成')
add_table(
    ['シーン', '想定場面'],
    [
        ['シーン1', '発災直後（初動対応）'],
        ['シーン2', '発災当日〜翌日（対応継続）'],
        ['シーン3', '複数日にわたる継続対応'],
    ],
    col_widths=[4, 11.5]
)

# ════════════════════════════════════════
# 2. 採点基準
# ════════════════════════════════════════
heading1('2. 採点基準・考え方（評価の仕様）')

heading2('評価コンピテンシー（5種）')
add_table(
    ['コンピテンシー', '評価する能力'],
    [
        ['コミュニケーション', '情報の共有・伝達・連携'],
        ['情報収集',           '状況把握・情報源の活用'],
        ['情報分析',           '収集情報の解釈・優先度判断'],
        ['想像・先読み',       'リスク予測・先行対応'],
        ['計画',               '行動計画・優先順位付け'],
    ],
    col_widths=[5, 10.5]
)
body('各問題には1〜3種のコンピテンシーが設定されており、問題ごとに異なります。', indent=True)

heading2('採点スケール')
add_table(
    ['点数', '基準'],
    [
        ['3点（優秀）',     '標準を明確に上回る対応'],
        ['2点（標準達成）', '合格水準（得点率計算の分母となる基準点）'],
        ['1点',             '部分的な達成'],
        ['0点',             '未達成'],
    ],
    col_widths=[4, 11.5]
)

heading2('得点率の計算')
body('分母：各問題のコンピテンシー数 × 2点（標準達成を100%の基準とする）')
body('実得点が分母を上回る場合、得点率は100%を超えることがあります。')
body('例：3コンピテンシーの問題で全て3点取得 → 9 / 6点（150%）', indent=True)

heading2('強制0点ルール')
body('以下に該当する場合、その問題のみ全コンピテンシーが強制0点（他の問題には影響なし）。レポートには警告バッジを表示。')
for item in [
    '① 空欄・無回答',
    '② 人命・安全を軽視した発言・考え方',
    '③ 自分だけを優先する自己中心的な考え方',
    '④ 他者に迷惑・過度な負担をかける行為',
]:
    bullet(item)

# ════════════════════════════════════════
# 3. レポートの構成
# ════════════════════════════════════════
heading1('3. レポートの構成')

heading2('冒頭あいさつ')
body('受験者名宛の挨拶・お礼文。得点率に応じた活躍期待の一文を含む。')
add_table(
    ['得点率', 'メッセージのトーン'],
    [
        ['100%以上', '組織を牽引するリーダーとしての活躍を期待'],
        ['80〜99%',  '知識・判断力をさらに磨いての活躍を期待'],
        ['60〜79%',  'さらなる研鑽を積んでの活躍を期待'],
        ['60%未満',  '今後の取り組みを通じた実力向上を期待'],
    ],
    col_widths=[4, 11.5]
)

heading2('① 受験者情報 ＋ 総合スコア')
for t in ['氏名・所属部署・ジャンル（班）', '総合得点と分母（例：119 / 102点）', '得点率（%）・コンピテンシー別達成率バーグラフ']:
    bullet(t)

heading2('② シーン別得点')
body('シーン1〜3ごとの得点・達成率・コンピテンシー内訳を表示。')

heading2('③ ジャンル内比較')
body('同班に複数受験者がいる場合のみ表示。ジャンル内順位・平均得点率・平均との差を表示。')

heading2('④ 総合評価')
add_table(
    ['段落', '内容'],
    [
        ['①', '得点率に応じた総評'],
        ['②', '達成率の高いコンピテンシーの称賛'],
        ['③', 'シーン別の振り返り'],
        ['④', '伸びしろのあるコンピテンシーへの改善提案'],
    ],
    col_widths=[2, 13.5]
)

heading2('⑤ 問題別採点結果（全22問）')
body('シーン単位でグループ化。各問について以下を表示：')
for t in [
    '問番号・問題種別・問題文',
    'コンピテンシー別得点（0〜3点、色分けバッジ）',
    '受験者の回答全文',
    '採点コメント・良かった点・不足点',
    '改善アドバイス',
]:
    bullet(t)

heading2('出力形式')
add_table(
    ['形式', '説明'],
    [
        ['Web画面',      'ブラウザで即時閲覧'],
        ['Word (.docx)', '印刷・編集可能なレポート'],
        ['PDF',          '150%初期ズーム設定。Webページと同等のデザイン（約20秒）'],
    ],
    col_widths=[4, 11.5]
)

# ════════════════════════════════════════
# 4. その他
# ════════════════════════════════════════
heading1('4. その他')

heading2('採点結果一覧')
for t in [
    '全受験者の得点・得点率・コンピテンシー別達成率を一覧表示',
    'ジャンル（班）別フィルタ',
    'ジャンル別サマリー（平均・最高・最低・コンピテンシー別平均）',
]:
    bullet(t)

heading2('技術仕様')
add_table(
    ['項目', '内容'],
    [
        ['採点AI',            'Claude Sonnet 4.6（Anthropic API）'],
        ['Webフレームワーク', 'Python / Flask'],
        ['PDF生成',           'Playwright（Chromium）によるHTMLレンダリング'],
        ['Word生成',          'python-docx'],
    ],
    col_widths=[4, 11.5]
)

heading2('制限事項')
for t in [
    'PDF生成には約20秒かかります（Chromiumのレイアウト処理）',
    'PDFダウンロードは通常のブラウザ（Chrome / Edge等）からご利用ください',
    '採点実行にはAnthropicのAPIキーが必要です',
]:
    bullet(t)

# ════════════════════════════════════════
# 補足：費用試算
# ════════════════════════════════════════
heading1('補足：費用試算（1人あたり）')

heading2('前提')
for t in [
    '採点API呼び出し：22回（1問1回）',
    '総合評価コメント：APIなし（テンプレート生成のためコスト0）',
    'システムプロンプトはキャッシュ機能により2問目以降は低コスト',
]:
    bullet(t)

heading2('トークン内訳（1問あたりの目安）')
add_table(
    ['要素', 'トークン数'],
    [
        ['システムプロンプト（キャッシュ）',     '約500'],
        ['前提条件・シーン説明',                 '約400'],
        ['問題文',                               '約150'],
        ['ルーブリック（1〜3コンピテンシー分）', '約300'],
        ['受験者の回答',                         '約200'],
        ['JSONテンプレート',                     '約100'],
        ['入力合計（1問）',                      '約1,150'],
        ['出力（スコア＋フィードバック）',       '約450'],
    ],
    col_widths=[8, 7.5]
)

heading2('22問合計のトークン数')
add_table(
    ['種別', '計算', 'トークン数'],
    [
        ['入力（ユーザープロンプト）',            '1,150 × 22問', '約25,300'],
        ['システムプロンプト キャッシュ書き込み', '500 × 1回',   '約500'],
        ['システムプロンプト キャッシュ読み出し', '500 × 21回',  '約10,500'],
        ['出力',                                  '450 × 22問',  '約9,900'],
    ],
    col_widths=[6.5, 4, 5]
)

heading2('費用計算（Claude Sonnet 4.6）')
add_table(
    ['種別', '単価', '費用'],
    [
        ['入力トークン',       '$3 / 100万',    '約$0.076'],
        ['キャッシュ書き込み', '$3.75 / 100万', '約$0.002'],
        ['キャッシュ読み出し', '$0.30 / 100万', '約$0.003'],
        ['出力トークン',       '$15 / 100万',   '約$0.149'],
        ['合計',               '',              '約$0.23（≒35円）'],
    ],
    col_widths=[6.5, 4, 5]
)

note('※ 1人あたり約$0.20〜$0.30（25〜45円）。回答が長い受験者ほど増加します。')
note('※ 8名全員を一括採点した場合の概算：約$1.6〜$2.5（250〜380円）')

out_path = 'bousai-grader/中上級者認定テスト_採点AIシステム仕様書.docx'
doc.save(out_path)
print('Saved:', out_path)
